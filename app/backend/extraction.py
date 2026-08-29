"""Extract plain text (with page/position tracking) from an uploaded CV.

DOCX and text-based PDF only for Phase 1 -- no OCR. A password-protected or
corrupted file raises ExtractionError with a human-readable message; callers
must not leak the underlying exception to the HR user.
"""
from dataclasses import dataclass
from pathlib import Path

import docx
from pypdf import PdfReader
from pypdf.errors import PdfReadError


class ExtractionError(Exception):
    """Raised with a message that is safe to show directly to an HR user."""


@dataclass
class TextBlock:
    text: str
    page: int | None  # None for DOCX (no native page concept)
    order: int


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
MC_NS = "http://schemas.openxmlformats.org/markup-compatibility/2006"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
W_P = f"{{{W_NS}}}p"
W_T = f"{{{W_NS}}}t"
W_TAB = f"{{{W_NS}}}tab"
W_BR = f"{{{W_NS}}}br"
W_HYPERLINK = f"{{{W_NS}}}hyperlink"
R_ID = f"{{{R_NS}}}id"
MC_FALLBACK = f"{{{MC_NS}}}Fallback"


def _own_text(paragraph) -> str:
    """Text belonging to this paragraph only, preserving its line structure.

    A paragraph can contain a drawing that contains further paragraphs (a
    text box). Collecting every descendant <w:t> would pull that nested
    content into the outer paragraph as well as emitting it in its own
    right, duplicating it.

    Line breaks and tabs are layout, but they are also word boundaries.
    Ignoring them runs neighbouring text together -- one CV produced
    "FIFTH GRADE TEACHERNorthwood Elementary School", two separate facts
    welded into a nonsense phrase, because a <w:br/> sat between them.
    """
    parts: list[str] = []
    for node in paragraph.iter():
        if node.tag not in (W_T, W_TAB, W_BR):
            continue
        ancestor = node.getparent()
        while ancestor is not None and ancestor.tag != W_P:
            ancestor = ancestor.getparent()
        if ancestor is not paragraph:
            continue
        if node.tag == W_T:
            parts.append(node.text or "")
        elif node.tag == W_TAB:
            parts.append(" ")
        else:  # a soft line break starts a new line, like any other
            parts.append("\n")
    return "".join(parts)


def _docx_paragraph_texts(document) -> list[str]:
    """Every paragraph in the document, in reading order.

    Walks the XML rather than using python-docx's `.paragraphs`/`.tables`,
    which only see the body: designed CV templates routinely lay their
    content out in TEXT BOXES, and that text is invisible to those APIs.
    One such template yielded 246 characters of cover-page blurb while the
    actual 3,800-character résumé sat in 52 text boxes.

    <mc:Fallback> branches are dropped first. Word stores a text box twice
    -- a modern DrawingML version and a legacy VML fallback carrying the
    same words -- so keeping both would duplicate the entire document.
    """
    root = document.element
    for fallback in list(root.iter(MC_FALLBACK)):
        parent = fallback.getparent()
        if parent is not None:
            parent.remove(fallback)

    seen: set[str] = set()
    texts: list[str] = []
    for paragraph in root.iter(W_P):
        # A paragraph may hold several visual lines (soft breaks); each is
        # its own logical line for classification purposes.
        for line in _own_text(paragraph).splitlines():
            text = line.strip()
            if not text:
                continue
            # Some templates repeat a heading in both a shape and the body.
            # Scoped to heading-SHAPED text only (short, and either ALL CAPS
            # or colon-terminated) -- a person's actual job title routinely
            # repeats verbatim between the letterhead and their most recent
            # job entry ("Marketing Operations Manager" at the top, then
            # again as the title of the current role), and deduping on raw
            # text alone silently deleted the second, genuine occurrence,
            # leaving that job entry with no title line at all.
            key = " ".join(text.split()).casefold()
            if key in seen and _is_heading_shaped(text):
                continue
            seen.add(key)
            texts.append(text)
    return texts


def _is_heading_shaped(text: str) -> bool:
    """Short and either ALL CAPS or colon-terminated -- a heading/label a
    template might duplicate between a decorative shape and the body flow,
    as opposed to ordinary CV content that can legitimately repeat."""
    stripped = text.strip()
    if len(stripped) > 60 or len(stripped.split()) > 8:
        return False
    if stripped.endswith(":"):
        return True
    letters = [c for c in stripped if c.isalpha()]
    if not letters:
        return False
    return sum(1 for c in letters if c.isupper()) / len(letters) >= 0.7


def _hyperlink_targets(document) -> list[str]:
    """URLs that exist only as Word hyperlink relationships.

    A CV that writes "ORCID profile" or "LinkedIn" as clickable text carries
    the actual address nowhere in the visible text -- it lives in
    word/_rels/document.xml.rels, keyed by the r:id on the <w:hyperlink>.
    Reading only <w:t> loses the identifier entirely, and identifiers are a
    section of the MDX template in their own right.

    Only targets whose display text does not already contain the URL are
    emitted, so a CV that spells its links out in full is not duplicated.
    """
    part = document.part
    out: list[str] = []
    seen: set[str] = set()
    for link in document.element.iter(W_HYPERLINK):
        rel_id = link.get(R_ID)
        if not rel_id:
            continue
        try:
            target = part.rels[rel_id].target_ref
        except KeyError:
            continue
        if not target or not target.lower().startswith(("http://", "https://", "mailto:")):
            continue
        shown = "".join(node.text or "" for node in link.iter(W_T))
        if target.rstrip("/") in shown.replace(" ", ""):
            continue
        # A mailto: whose address is already the visible text carries nothing
        # new -- the email is extracted from the text in the normal way, and
        # emitting "name@x.ac.ae: mailto:name@x.ac.ae" adds only noise.
        if target.lower().startswith("mailto:") and target[7:].strip() in shown:
            continue
        key = target.rstrip("/").casefold()
        if key in seen:
            continue
        seen.add(key)
        label = " ".join(shown.split())
        out.append(f"{label}: {target}" if label else target)
    return out


def _docx_header_footer_texts(document) -> list[str]:
    """Text from headers and footers of every section of the document.

    These are separate parts of the package, so walking word/document.xml
    never sees them. CVs commonly put the contact strip, ORCID, or a running
    name/page line there, and §3 of the conversion spec requires the whole
    document be read regardless of where content sits.
    """
    texts: list[str] = []
    for section in document.sections:
        for container in (
            section.header, section.first_page_header, section.even_page_header,
            section.footer, section.first_page_footer, section.even_page_footer,
        ):
            try:
                element = container._element
            except Exception:
                continue
            for paragraph in element.iter(W_P):
                for line in _own_text(paragraph).splitlines():
                    text = line.strip()
                    if text:
                        texts.append(text)
    return texts


def extract_docx(path: Path) -> list[TextBlock]:
    try:
        document = docx.Document(str(path))
    except Exception as exc:  # python-docx raises bare Exception/KeyError on bad zips
        raise ExtractionError(
            "This DOCX file could not be opened. It may be corrupted or not a "
            "valid Word document."
        ) from exc

    blocks: list[TextBlock] = []
    order = 0

    # Hyperlink targets are collected before the body walk strips mc:Fallback
    # branches, and emitted first so a URL-only line is never mistaken for a
    # continuation of the paragraph above it.
    try:
        link_lines = _hyperlink_targets(document)
    except Exception:
        link_lines = []  # a malformed rels part must not fail the whole upload

    try:
        header_lines = _docx_header_footer_texts(document)
    except Exception:
        header_lines = []

    seen_keys: set[str] = set()
    for text in _docx_paragraph_texts(document) + header_lines + link_lines:
        key = " ".join(text.split()).casefold()
        # Same reasoning as the body-level dedup above: a running header or
        # footer commonly repeats a short label (name, page strip) that IS a
        # genuine duplicate, but content-shaped text (a job title, a real
        # sentence) is not deduplicated against just because it happens to
        # match something seen earlier in the body.
        if key in seen_keys and _is_heading_shaped(text):
            continue
        seen_keys.add(key)
        blocks.append(TextBlock(text=text, page=None, order=order))
        order += 1

    if not blocks:
        raise ExtractionError(
            "No readable text was found in this document. It may be empty or "
            "contain only images -- image-only CVs require OCR, which isn't "
            "supported yet."
        )
    return blocks


def extract_pdf(path: Path) -> list[TextBlock]:
    try:
        reader = PdfReader(str(path))
    except PdfReadError as exc:
        raise ExtractionError(
            "This PDF could not be opened. It may be corrupted."
        ) from exc

    if reader.is_encrypted:
        try:
            result = reader.decrypt("")
            if result == 0:
                raise ExtractionError(
                    "This PDF is password-protected. Please remove the password "
                    "and upload it again."
                )
        except Exception as exc:
            raise ExtractionError(
                "This PDF is password-protected. Please remove the password "
                "and upload it again."
            ) from exc

    blocks: list[TextBlock] = []
    order = 0
    for page_num, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        if text:
            blocks.append(TextBlock(text=text, page=page_num, order=order))
            order += 1

    if not blocks:
        raise ExtractionError(
            "No selectable text was found in this PDF. It looks like a scanned "
            "or image-based document -- OCR support isn't available yet, so "
            "this file can't be processed automatically."
        )
    return blocks


import re

# Share of whitespace-separated tokens that are a single stray letter.
# Documents built on a subsetted font with a broken character map extract as
# text but not as WORDS: "NATIONALITY" comes out as "N AT I z N A 9 I T B",
# "Orlando" as "zrlando", "example@email.com" as "ePample/email.com". The
# file is not empty and not a scan, so neither existing check catches it, and
# every downstream stage then works confidently on nonsense.
#
# Measured across the corpus: two such files score 0.21-0.22, and the highest
# legitimate CV scores 0.049 -- a design-heavy résumé whose letter-spaced
# headings are real text. The threshold sits with a wide margin either side.
GARBLED_SINGLE_LETTER_RATIO = 0.12
MIN_TOKENS_FOR_GARBLE_CHECK = 100


def _looks_garbled(text: str) -> bool:
    tokens = text.split()
    if len(tokens) < MIN_TOKENS_FOR_GARBLE_CHECK:
        return False
    strays = sum(1 for t in tokens if len(t) == 1 and t.isalpha())
    return strays / len(tokens) > GARBLED_SINGLE_LETTER_RATIO


def extract(path: Path) -> list[TextBlock]:
    suffix = path.suffix.lower()
    if suffix == ".docx":
        blocks = extract_docx(path)
    elif suffix == ".pdf":
        blocks = extract_pdf(path)
    else:
        raise ExtractionError(f"Unsupported file type '{suffix}'. Please upload a DOCX or PDF.")

    if _looks_garbled(blocks_to_plain_text(blocks)):
        raise ExtractionError(
            "The text in this file is scrambled and cannot be read reliably -- "
            "words come out broken into single letters, and names, emails and "
            "dates are corrupted. This usually means the document was produced "
            "by a converter that embedded a damaged font. Please re-save it as "
            "a PDF from the original application, or upload the original file."
        )
    return blocks


def blocks_to_plain_text(blocks: list[TextBlock]) -> str:
    """Join blocks into one text stream, one block per line, for the classifier."""
    return "\n".join(b.text for b in blocks)
